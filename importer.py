import docx
import json

def from_docx(path):
    doc = docx.Document(path)
    p_count = 0
    quiz_count = 0
    quest_count = 0
    var_count = -1
    quiz = {}
    for element in doc.paragraphs:
        #print(element.text, "\n")
        try:
            if element.text[0].isupper():
                if doc.paragraphs[p_count+1].text[0].isupper():
                    quiz_count+=1
                    quiz[f"quiz{quiz_count}"] = {}
                    quiz[f"quiz{quiz_count}"]["name"] = element.text
                    quiz[f"quiz{quiz_count}"]["questions"] = {}
                else:
                    quest_count+=1
                    quiz[f"quiz{quiz_count}"]["questions"][f"question{quest_count}"] = {}
                    quiz[f"quiz{quiz_count}"]["questions"][f"question{quest_count}"]["question"] = element.text
                    quiz[f"quiz{quiz_count}"]["questions"][f"question{quest_count}"]["variants"] = {}
                    var_count=-1
                    #print(element.text)
            elif element.text[0].islower() or element.text[0].isdigit():
                var_count+=1
                quiz[f"quiz{quiz_count}"]["questions"][f"question{quest_count}"]["variants"][f"variant{var_count}"] = element.text
                if element.runs[0].font.bold == True:
                    quiz[f"quiz{quiz_count}"]["questions"][f"question{quest_count}"]["variants"]["correct"] = var_count
                #print(var_count, element.text, "\n")
            else:
                pass
        except Exception as e:
            print(e)
            print("НЕ СТАВЬ ЁБАНЫЕ ПРОБЕЛЫ МЕЖДУ ПАРАГРАФАМИ, ЭТО НЕ ДЛЯ ЛЮДЕЙ НАПИСАНО")
        p_count+=1
    #print(json.dumps(quiz, sort_keys=True, indent=4, ensure_ascii=False))
    return(quiz)

        # print(element._p.xml)
        # root = ET.fromstring(element._p.xml)
        # for child in root:
        #     print(child.tag, child.attrib)
        #     for elem in child:
        #         print(elem.text)



if __name__=="__main__":
    from_docx("./import/doc.docx")